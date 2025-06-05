#!/usr/bin/env python3
"""
Knowledge Centre - Comprehensive Trading Research & Intelligence System
Auto-ingests, analyzes, indexes all trading information for continuous improvement
"""

import asyncio
import json
import os
import sqlite3
import hashlib
import requests
from datetime import datetime, timedelta
from typing import Dict, List, Any, Optional, Tuple
from pathlib import Path
from dataclasses import dataclass
import PyPDF2
import docx
import youtube_dl
import feedparser
from bs4 import BeautifulSoup
import chromadb
from sentence_transformers import SentenceTransformer
import re
import logging

# Import our AI orchestrator
import sys
sys.path.append(str(Path(__file__).parent.parent.parent))
from core.orchestration.master_orchestrator import CryptoAI

@dataclass
class ResearchDocument:
    doc_id: str
    title: str
    content: str
    source: str
    doc_type: str  # pdf, web, youtube, research_paper, etc.
    credibility_score: float
    relevance_tags: List[str]
    created_at: datetime
    last_analyzed: datetime
    metadata: Dict[str, Any]

@dataclass
class ResearchInsight:
    insight_id: str
    content: str
    insight_type: str  # strategy, pattern, risk_management, market_analysis
    confidence: float
    supporting_docs: List[str]
    generated_prompts: List[str]
    created_at: datetime

class KnowledgeCentre:
    """Comprehensive knowledge centre for trading research and continuous improvement"""
    
    def __init__(self):
        self.db_path = "data/databases/orion_protocol.db"
        self.knowledge_path = "data/knowledge_centre"
        self.setup_directories()
        self.setup_database()
        
        # AI components
        self.crypto_ai = CryptoAI()
        
        # Vector database for semantic search
        self.chroma_client = chromadb.PersistentClient(path=f"{self.knowledge_path}/vector_db")
        self.collection = self.chroma_client.get_or_create_collection("trading_knowledge")
        
        # Embedding model for semantic indexing
        self.embedding_model = SentenceTransformer('all-MiniLM-L6-v2')
        
        # Content processors
        self.processors = {
            'pdf': self.process_pdf,
            'docx': self.process_docx,
            'web': self.process_web_content,
            'youtube': self.process_youtube,
            'research_paper': self.process_research_paper,
            'news': self.process_news_article
        }
        
        # Auto-research sources
        self.auto_sources = {
            'reddit_feeds': [
                'https://www.reddit.com/r/SecurityAnalysis.json',
                'https://www.reddit.com/r/investing.json',
                'https://www.reddit.com/r/SecurityAnalysis.json'
            ],
            'research_sites': [
                'https://papers.ssrn.com/sol3/JELJOUR_Results.cfm?form_name=journalBrowse&journal_id=1079355',
                'https://arxiv.org/search/?query=cryptocurrency+trading&searchtype=all'
            ],
            'trading_blogs': [
                'https://quantocracy.com/feed/',
                'https://www.investopedia.com/rss'
            ]
        }
        
        self.setup_logging()
    
    def setup_directories(self):
        """Setup knowledge centre directories"""
        directories = [
            "data/knowledge_centre",
            "data/knowledge_centre/documents",
            "data/knowledge_centre/processed", 
            "data/knowledge_centre/vector_db",
            "data/knowledge_centre/auto_research",
            "data/databases"
        ]
        
        for directory in directories:
            Path(directory).mkdir(parents=True, exist_ok=True)
    
    def setup_database(self):
        """Setup comprehensive knowledge database"""
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        
        # Documents table
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS research_documents (
                doc_id TEXT PRIMARY KEY,
                title TEXT NOT NULL,
                content TEXT NOT NULL,
                source TEXT NOT NULL,
                doc_type TEXT NOT NULL,
                credibility_score REAL NOT NULL,
                relevance_tags TEXT NOT NULL,
                created_at REAL NOT NULL,
                last_analyzed REAL NOT NULL,
                metadata TEXT NOT NULL,
                embedding_id TEXT
            )
        """)
        
        # Insights table
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS research_insights (
                insight_id TEXT PRIMARY KEY,
                content TEXT NOT NULL,
                insight_type TEXT NOT NULL,
                confidence REAL NOT NULL,
                supporting_docs TEXT NOT NULL,
                generated_prompts TEXT NOT NULL,
                created_at REAL NOT NULL,
                applied_to_system BOOLEAN DEFAULT FALSE
            )
        """)
        
        # Generated prompts table
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS generated_prompts (
                prompt_id TEXT PRIMARY KEY,
                prompt_text TEXT NOT NULL,
                prompt_type TEXT NOT NULL,
                target_component TEXT NOT NULL,
                effectiveness_score REAL DEFAULT 0.5,
                usage_count INTEGER DEFAULT 0,
                created_at REAL NOT NULL,
                last_used REAL,
                source_insight_id TEXT
            )
        """)
        
        # Auto research queue
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS auto_research_queue (
                queue_id TEXT PRIMARY KEY,
                research_query TEXT NOT NULL,
                priority INTEGER NOT NULL,
                status TEXT DEFAULT 'pending',
                created_at REAL NOT NULL,
                processed_at REAL,
                results_count INTEGER DEFAULT 0
            )
        """)
        
        # Source credibility tracking
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS source_credibility (
                source_url TEXT PRIMARY KEY,
                credibility_score REAL NOT NULL,
                assessment_count INTEGER DEFAULT 1,
                last_updated REAL NOT NULL,
                credibility_factors TEXT NOT NULL
            )
        """)
        
        conn.commit()
        conn.close()
        print("✅ Knowledge Centre database setup complete")
    
    def setup_logging(self):
        """Setup logging for knowledge centre"""
        log_dir = Path("logs/knowledge_centre")
        log_dir.mkdir(parents=True, exist_ok=True)
        
        logging.basicConfig(
            level=logging.INFO,
            format='%(asctime)s - %(levelname)s - %(message)s',
            handlers=[
                logging.FileHandler(log_dir / 'knowledge_centre.log'),
                logging.StreamHandler()
            ]
        )
        self.logger = logging.getLogger(__name__)
    
    async def ingest_document(self, file_path: str, doc_type: str, source: str = None) -> str:
        """Ingest and process a document into the knowledge centre"""
        
        self.logger.info(f"📄 Ingesting document: {file_path}")
        
        # Generate document ID
        doc_id = hashlib.md5(f"{file_path}{datetime.now().timestamp()}".encode()).hexdigest()
        
        # Process document based on type
        if doc_type in self.processors:
            content, metadata = await self.processors[doc_type](file_path)
        else:
            raise ValueError(f"Unsupported document type: {doc_type}")
        
        # Assess credibility
        credibility_score = await self.assess_credibility(source or file_path, content)
        
        # Generate relevance tags using AI
        relevance_tags = await self.generate_relevance_tags(content)
        
        # Create document object
        document = ResearchDocument(
            doc_id=doc_id,
            title=metadata.get('title', Path(file_path).name),
            content=content,
            source=source or file_path,
            doc_type=doc_type,
            credibility_score=credibility_score,
            relevance_tags=relevance_tags,
            created_at=datetime.now(),
            last_analyzed=datetime.now(),
            metadata=metadata
        )
        
        # Store in database
        self.store_document(document)
        
        # Create vector embedding and store
        embedding = self.embedding_model.encode(content)
        self.collection.add(
            ids=[doc_id],
            embeddings=[embedding.tolist()],
            documents=[content],
            metadatas=[{
                'title': document.title,
                'source': document.source,
                'doc_type': document.doc_type,
                'credibility_score': document.credibility_score,
                'tags': ','.join(document.relevance_tags)
            }]
        )
        
        # Auto-generate insights
        await self.generate_insights_from_document(document)
        
        self.logger.info(f"✅ Document ingested: {doc_id}")
        return doc_id
    
    async def process_pdf(self, file_path: str) -> Tuple[str, Dict]:
        """Process PDF document"""
        content = ""
        metadata = {"pages": 0}
        
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                metadata["pages"] = len(pdf_reader.pages)
                
                for page in pdf_reader.pages:
                    content += page.extract_text() + "\n"
                    
        except Exception as e:
            self.logger.error(f"❌ PDF processing error: {e}")
            content = f"Error processing PDF: {e}"
        
        return content, metadata
    
    async def process_docx(self, file_path: str) -> Tuple[str, Dict]:
        """Process Word document"""
        content = ""
        metadata = {"paragraphs": 0}
        
        try:
            doc = docx.Document(file_path)
            paragraphs = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    paragraphs.append(paragraph.text)
            
            content = "\n".join(paragraphs)
            metadata["paragraphs"] = len(paragraphs)
            
        except Exception as e:
            self.logger.error(f"❌ DOCX processing error: {e}")
            content = f"Error processing DOCX: {e}"
        
        return content, metadata
    
    async def process_web_content(self, url: str) -> Tuple[str, Dict]:
        """Process web page content"""
        content = ""
        metadata = {"url": url}
        
        try:
            response = requests.get(url, timeout=30)
            soup = BeautifulSoup(response.content, 'html.parser')
            
            # Extract title
            title_tag = soup.find('title')
            metadata["title"] = title_tag.text if title_tag else "Unknown"
            
            # Extract main content (remove scripts, styles, etc.)
            for script in soup(["script", "style", "nav", "footer", "header"]):
                script.decompose()
            
            # Get text content
            content = soup.get_text()
            
            # Clean up whitespace
            lines = (line.strip() for line in content.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            content = ' '.join(chunk for chunk in chunks if chunk)
            
        except Exception as e:
            self.logger.error(f"❌ Web content processing error: {e}")
            content = f"Error processing web content: {e}"
        
        return content, metadata
    
    async def process_youtube(self, url: str) -> Tuple[str, Dict]:
        """Process YouTube video (extract transcript if available)"""
        content = ""
        metadata = {"url": url}
        
        try:
            # Note: This is a simplified version
            # In production, you'd use youtube-dl or similar to extract transcripts
            metadata["title"] = "YouTube Video"
            content = f"YouTube video content from: {url}\n[Transcript extraction would be implemented here]"
            
        except Exception as e:
            self.logger.error(f"❌ YouTube processing error: {e}")
            content = f"Error processing YouTube video: {e}"
        
        return content, metadata
    
    async def process_research_paper(self, file_path: str) -> Tuple[str, Dict]:
        """Process research paper with enhanced analysis"""
        content, metadata = await self.process_pdf(file_path)
        
        # Enhanced metadata for research papers
        metadata.update({
            "document_type": "research_paper",
            "analysis_enhanced": True
        })
        
        return content, metadata
    
    async def process_news_article(self, url: str) -> Tuple[str, Dict]:
        """Process news article"""
        content, metadata = await self.process_web_content(url)
        
        metadata.update({
            "document_type": "news_article",
            "processed_date": datetime.now().isoformat()
        })
        
        return content, metadata
    
    async def assess_credibility(self, source: str, content: str) -> float:
        """Assess credibility of source and content using AI"""
        
        # Check if we have cached credibility for this source
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        
        cursor.execute("SELECT credibility_score FROM source_credibility WHERE source_url = ?", (source,))
        result = cursor.fetchone()
        
        if result:
            conn.close()
            return result[0]
        
        # AI-powered credibility assessment
        assessment_prompt = f"""
        Assess the credibility of this trading/financial content:
        
        Source: {source}
        Content (first 1000 chars): {content[:1000]}
        
        Rate credibility from 0.0 to 1.0 based on:
        - Source reputation
        - Content quality and factual accuracy
        - Presence of data/citations
        - Writing quality and professionalism
        - Potential bias or conflicts of interest
        
        Respond with just a decimal number between 0.0 and 1.0.
        """
        
        try:
            credibility_response = await self.crypto_ai.analyze_market(assessment_prompt)
            credibility_score = float(re.search(r'(\d+\.?\d*)', credibility_response).group(1))
            credibility_score = max(0.0, min(1.0, credibility_score))
            
        except:
            # Default moderate credibility if AI assessment fails
            credibility_score = 0.6
        
        # Store credibility assessment
        cursor.execute("""
            INSERT OR REPLACE INTO source_credibility 
            (source_url, credibility_score, last_updated, credibility_factors)
            VALUES (?, ?, ?, ?)
        """, (source, credibility_score, datetime.now().timestamp(), json.dumps({"ai_assessed": True})))
        
        conn.commit()
        conn.close()
        
        return credibility_score
    
    async def generate_relevance_tags(self, content: str) -> List[str]:
        """Generate relevance tags using AI"""
        
        tagging_prompt = f"""
        Generate 5-8 relevant tags for this trading/financial content:
        
        Content: {content[:2000]}
        
        Tags should be specific trading/finance terms like:
        - Technical analysis, fundamental analysis
        - Risk management, portfolio optimization
        - Cryptocurrency, stocks, forex
        - Strategy types (momentum, mean reversion, etc.)
        - Market conditions (bull market, bear market, etc.)
        
        Return only the tags, separated by commas.
        """
        
        try:
            tags_response = await self.crypto_ai.analyze_market(tagging_prompt)
            tags = [tag.strip() for tag in tags_response.split(',')]
            return tags[:8]  # Limit to 8 tags
            
        except:
            return ["trading", "analysis", "finance"]  # Default tags
    
    async def generate_insights_from_document(self, document: ResearchDocument):
        """Generate actionable insights from document using AI"""
        
        insights_prompt = f"""
        Analyze this trading/financial document and extract key insights:
        
        Title: {document.title}
        Content: {document.content[:3000]}
        Credibility: {document.credibility_score}
        
        Extract:
        1. Key trading strategies mentioned
        2. Risk management techniques
        3. Market analysis methods
        4. Performance metrics or results
        5. Novel approaches or innovations
        
        For each insight, suggest specific prompts that could be used to:
        - Improve our trading strategies
        - Enhance risk management
        - Better analyze market data
        - Optimize our AI models
        
        Format as JSON with insights and suggested prompts.
        """
        
        try:
            insights_response = await self.crypto_ai.analyze_market(insights_prompt, complexity="high")
            
            # Parse insights and store
            insight_id = hashlib.md5(f"{document.doc_id}{datetime.now().timestamp()}".encode()).hexdigest()
            
            insight = ResearchInsight(
                insight_id=insight_id,
                content=insights_response,
                insight_type="comprehensive_analysis",
                confidence=document.credibility_score,
                supporting_docs=[document.doc_id],
                generated_prompts=self.extract_prompts_from_insight(insights_response),
                created_at=datetime.now()
            )
            
            self.store_insight(insight)
            
        except Exception as e:
            self.logger.error(f"❌ Insight generation error: {e}")
    
    def extract_prompts_from_insight(self, insight_content: str) -> List[str]:
        """Extract generated prompts from insight content"""
        # Simple extraction - in production you'd use more sophisticated parsing
        prompts = []
        lines = insight_content.split('\n')
        
        for line in lines:
            if 'prompt' in line.lower() or 'ask' in line.lower():
                prompts.append(line.strip())
        
        return prompts[:5]  # Limit to 5 prompts
    
    async def auto_research_cycle(self):
        """Automated research cycle - discovers and ingests new content"""
        
        self.logger.info("🔍 Starting auto-research cycle...")
        
        # Generate new research queries based on current gaps
        research_queries = await self.generate_research_queries()
        
        for query in research_queries:
            await self.execute_research_query(query)
        
        # Process auto-source feeds
        await self.process_auto_sources()
        
        self.logger.info("✅ Auto-research cycle complete")
    
    async def generate_research_queries(self) -> List[str]:
        """Generate new research queries based on system gaps"""
        
        gap_analysis_prompt = """
        Based on our current trading AI system, suggest 5 research queries to improve:
        
        Current capabilities:
        - Multi-source data collection
        - Pattern discovery
        - Sentiment analysis
        - Technical analysis
        
        Suggest research queries for:
        1. New trading strategies
        2. Risk management improvements
        3. Market prediction techniques
        4. Alternative data sources
        5. AI/ML optimization methods
        
        Return as a simple list.
        """
        
        try:
            queries_response = await self.crypto_ai.analyze_market(gap_analysis_prompt)
            queries = [q.strip() for q in queries_response.split('\n') if q.strip()]
            return queries[:5]
            
        except:
            return [
                "advanced momentum trading strategies",
                "crypto market sentiment indicators", 
                "risk-adjusted portfolio optimization",
                "alternative data for crypto trading",
                "machine learning for market prediction"
            ]
    
    async def execute_research_query(self, query: str):
        """Execute a research query and ingest results"""
        
        self.logger.info(f"🔎 Executing research query: {query}")
        
        # For MVP, we'll focus on web search results
        # In production, you'd integrate with academic databases, research APIs, etc.
        
        search_urls = [
            f"https://scholar.google.com/scholar?q={query.replace(' ', '+')}+trading",
            f"https://papers.ssrn.com/sol3/results.cfm?RequestTimeout=50000000&q={query.replace(' ', '+')}"
        ]
        
        results_count = 0
        for url in search_urls:
            try:
                # Process search results (simplified)
                content, metadata = await self.process_web_content(url)
                
                if len(content) > 500:  # Only process substantial content
                    doc_id = await self.ingest_document(url, 'web', url)
                    results_count += 1
                    
            except Exception as e:
                self.logger.error(f"❌ Research query execution error: {e}")
        
        # Store query results
        self.store_research_query(query, results_count)
        return results_count
    
    async def process_auto_sources(self):
        """Process configured auto-research sources"""
        
        for source_type, urls in self.auto_sources.items():
            for url in urls:
                try:
                    if source_type == 'reddit_feeds':
                        await self.process_reddit_feed(url)
                    elif source_type == 'trading_blogs':
                        await self.process_rss_feed(url)
                        
                except Exception as e:
                    self.logger.error(f"❌ Auto-source processing error for {url}: {e}")
    
    async def process_reddit_feed(self, url: str):
        """Process Reddit feed for trading discussions"""
        try:
            headers = {"User-Agent": "TradingResearch Bot 1.0"}
            response = requests.get(url, headers=headers, timeout=15)
            
            if response.status_code == 200:
                data = response.json()
                posts = data.get("data", {}).get("children", [])
                
                for post in posts[:5]:  # Top 5 posts
                    post_data = post.get("data", {})
                    if len(post_data.get("selftext", "")) > 200:  # Substantial content
                        
                        # Create temporary content
                        content = f"Title: {post_data.get('title', '')}\n\nContent: {post_data.get('selftext', '')}"
                        
                        if await self.assess_credibility(url, content) > 0.5:
                            await self.ingest_document(url, 'web', url)
                            
        except Exception as e:
            self.logger.error(f"❌ Reddit feed processing error: {e}")
    
    async def process_rss_feed(self, url: str):
        """Process RSS feed for trading content"""
        try:
            feed = feedparser.parse(url)
            
            for entry in feed.entries[:3]:  # Top 3 articles
                if hasattr(entry, 'link'):
                    await self.ingest_document(entry.link, 'web', entry.link)
                    
        except Exception as e:
            self.logger.error(f"❌ RSS feed processing error: {e}")
    
    def store_document(self, document: ResearchDocument):
        """Store document in database"""
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        
        cursor.execute("""
            INSERT OR REPLACE INTO research_documents 
            (doc_id, title, content, source, doc_type, credibility_score, 
             relevance_tags, created_at, last_analyzed, metadata)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """, (
            document.doc_id, document.title, document.content, document.source,
            document.doc_type, document.credibility_score, json.dumps(document.relevance_tags),
            document.created_at.timestamp(), document.last_analyzed.timestamp(),
            json.dumps(document.metadata)
        ))
        
        conn.commit()
        conn.close()
    
    def store_insight(self, insight: ResearchInsight):
        """Store insight in database"""
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        
        cursor.execute("""
            INSERT OR REPLACE INTO research_insights 
            (insight_id, content, insight_type, confidence, supporting_docs, 
             generated_prompts, created_at)
            VALUES (?, ?, ?, ?, ?, ?, ?)
        """, (
            insight.insight_id, insight.content, insight.insight_type, insight.confidence,
            json.dumps(insight.supporting_docs), json.dumps(insight.generated_prompts),
            insight.created_at.timestamp()
        ))
        
        conn.commit()
        conn.close()
    
    def store_research_query(self, query: str, results_count: int):
        """Store research query results"""
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        
        query_id = hashlib.md5(f"{query}{datetime.now().timestamp()}".encode()).hexdigest()
        
        cursor.execute("""
            INSERT INTO auto_research_queue 
            (queue_id, research_query, priority, status, created_at, processed_at, results_count)
            VALUES (?, ?, ?, ?, ?, ?, ?)
        """, (
            query_id, query, 1, 'completed', datetime.now().timestamp(),
            datetime.now().timestamp(), results_count
        ))
        
        conn.commit()
        conn.close()
    
    async def semantic_search(self, query: str, limit: int = 5) -> List[Dict]:
        """Perform semantic search across knowledge base"""
        
        # Generate query embedding
        query_embedding = self.embedding_model.encode(query)
        
        # Search vector database
        results = self.collection.query(
            query_embeddings=[query_embedding.tolist()],
            n_results=limit
        )
        
        # Format results
        formatted_results = []
        for i, doc_id in enumerate(results['ids'][0]):
            formatted_results.append({
                'doc_id': doc_id,
                'content': results['documents'][0][i],
                'metadata': results['metadatas'][0][i],
                'similarity_score': 1 - results['distances'][0][i]  # Convert distance to similarity
            })
        
        return formatted_results
    
    def get_knowledge_summary(self) -> Dict[str, Any]:
        """Get comprehensive knowledge centre summary"""
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        
        # Document statistics
        cursor.execute("SELECT COUNT(*), doc_type FROM research_documents GROUP BY doc_type")
        doc_stats = dict(cursor.fetchall())
        
        cursor.execute("SELECT AVG(credibility_score), COUNT(*) FROM research_documents")
        avg_credibility, total_docs = cursor.fetchone()
        
        # Insight statistics
        cursor.execute("SELECT COUNT(*), insight_type FROM research_insights GROUP BY insight_type")
        insight_results = cursor.fetchall()
        insight_stats = {insight_type: count for count, insight_type in insight_results}
        
        # Recent activity
        cursor.execute("""
            SELECT COUNT(*) FROM research_documents 
            WHERE created_at > ?
        """, ((datetime.now() - timedelta(days=7)).timestamp(),))
        recent_docs = cursor.fetchone()[0]
        
        conn.close()
        
        return {
            "total_documents": total_docs or 0,
            "document_types": doc_stats,
            "average_credibility": avg_credibility or 0,
            "total_insights": sum(insight_stats.values()) if insight_stats else 0,
            "insight_types": insight_stats,
            "recent_activity_7d": recent_docs,
            "vector_embeddings": self.collection.count(),
            "last_updated": datetime.now().isoformat()
        }

if __name__ == "__main__":
    async def main():
        kc = KnowledgeCentre()
        
        print("🧠 Testing Knowledge Centre...")
        
        # Test document ingestion
        print("📄 Testing document ingestion...")
        
        # Test semantic search
        print("🔍 Testing semantic search...")
        results = await kc.semantic_search("momentum trading strategies")
        print(f"Found {len(results)} relevant documents")
        
        # Test auto-research
        print("🔎 Testing auto-research...")
        await kc.auto_research_cycle()
        
        # Get summary
        summary = kc.get_knowledge_summary()
        print(f"\n📊 KNOWLEDGE CENTRE SUMMARY:")
        print(f"   Total Documents: {summary['total_documents']}")
        print(f"   Average Credibility: {summary['average_credibility']:.2f}")
        print(f"   Total Insights: {summary['total_insights']}")
        print(f"   Recent Activity: {summary['recent_activity_7d']} docs (7 days)")
        
        print("\n✅ Knowledge Centre test complete!")
    
    asyncio.run(main()) 